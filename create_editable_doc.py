"""
Create Editable Word Document from Instructions
Creates a .docx file that can be edited in Microsoft Word
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    def create_word_document():
        """Create editable Word document from customer instructions"""
        
        input_file = "CUSTOMER_DOWNLOAD_INSTRUCTIONS.txt"
        output_file = "CUSTOMER_DOWNLOAD_INSTRUCTIONS.docx"
        
        if not os.path.exists(input_file):
            print(f"❌ File not found: {input_file}")
            return
        
        print("\n" + "="*70)
        print("📄 CREATING EDITABLE WORD DOCUMENT")
        print("="*70)
        print()
        
        # Read the text file
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        print("⏳ Generating Word document...")
        
        for line in lines:
            line = line.rstrip()
            
            # Skip separator lines
            if line.strip().startswith('====') or line.strip().startswith('────'):
                continue
            
            # Empty lines
            if not line.strip():
                doc.add_paragraph()
                continue
            
            # Main title
            if '🎃 HALLOWEEN CHARACTER BUNDLE' in line:
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 107, 0)  # Orange
                continue
            
            # Section headings (with emojis)
            if any(emoji in line for emoji in ['📥', '⚠️', '❓', '📞', '🎨', '✨', '💡', '🌟']):
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(51, 51, 51)
                continue
            
            # Step headings
            if line.strip().startswith('STEP '):
                p = doc.add_paragraph(line.strip())
                p.runs[0].font.size = Pt(12)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 107, 0)  # Orange
                continue
            
            # Problem headings
            if line.strip().startswith('PROBLEM:'):
                p = doc.add_paragraph(line.strip())
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(200, 0, 0)  # Red
                continue
            
            # Solution/Note labels
            if line.strip().startswith(('SOLUTION:', 'NOTE:', 'METHOD', 'INCLUDE IN')):
                p = doc.add_paragraph(line.strip())
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.bold = True
                continue
            
            # Regular text (including Google Drive links)
            if line.strip():
                text = line.strip()
                
                # Check if this is the Google Drive link line
                if 'https://drive.google.com' in text:
                    p = doc.add_paragraph()
                    
                    # Split at link to make it clickable
                    if '👉' in text:
                        before_link = text.split('https://')[0]
                        link_part = 'https://' + text.split('https://')[1].split('👈')[0].strip()
                        after_link = '👈' if '👈' in text else ''
                        
                        # Add text before link
                        run1 = p.add_run(before_link)
                        run1.font.size = Pt(10)
                        
                        # Add hyperlink
                        run2 = p.add_run(link_part)
                        run2.font.size = Pt(10)
                        run2.font.color.rgb = RGBColor(0, 102, 204)  # Blue
                        run2.font.underline = True
                        
                        # Add text after link
                        if after_link:
                            run3 = p.add_run(' ' + after_link)
                            run3.font.size = Pt(10)
                    else:
                        # Just add the whole line
                        run = p.add_run(text)
                        run.font.size = Pt(10)
                        if 'https://' in text:
                            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
                            run.font.underline = True
                else:
                    # Regular paragraph
                    p = doc.add_paragraph(text)
                    p.runs[0].font.size = Pt(10)
        
        # Save Word document
        doc.save(output_file)
        
        # Get file size
        file_size = os.path.getsize(output_file) / 1024  # KB
        
        print("\n" + "="*70)
        print("✅ WORD DOCUMENT CREATED SUCCESSFULLY!")
        print("="*70)
        print(f"📁 File: {output_file}")
        print(f"📊 Size: {file_size:.2f} KB")
        print()
        print("✏️ This document is FULLY EDITABLE in Microsoft Word!")
        print("📝 You can:")
        print("   - Edit any text")
        print("   - Change colors and fonts")
        print("   - Add your branding")
        print("   - Export as PDF when done")
        print()
        print("💡 To make PDF from Word:")
        print("   File → Save As → PDF")
        print("="*70)
    
    if __name__ == "__main__":
        create_word_document()
        
except ImportError:
    print("\n" + "="*70)
    print("📦 INSTALLING PYTHON-DOCX")
    print("="*70)
    print()
    print("Installing python-docx library...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    print("\n✅ Installed! Running script again...")
    print("="*70)
    
    # Import again and run
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    
    # Paste the entire create_word_document function here again
    print("\nPlease run the script again: python create_editable_doc.py")
